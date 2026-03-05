from docx import Document
from docx.shared import Pt
from datetime import datetime
import os

output_dir = r"e:\toolCode\AuthService\docs"
os.makedirs(output_dir, exist_ok=True)
file_path = os.path.join(output_dir, "Tai-lieu-he-thong-SSO-AuthService.docx")

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

doc.add_heading('TÀI LIỆU HỆ THỐNG SSO - .NET 10 (API + MVC + SQL Server)', 0)
doc.add_paragraph(f'Ngày cập nhật: {datetime.now().strftime("%d/%m/%Y %H:%M")}')
doc.add_paragraph('Phạm vi: Hệ thống gồm AuthService (API), Application1 (MVC), Application2 (MVC) sử dụng xác thực tập trung SSO bằng JWT và Cookie dùng chung.')

doc.add_heading('1. Tổng quan hệ thống', level=1)
doc.add_paragraph('Mục tiêu: xây dựng hệ thống đăng nhập tập trung (Single Sign-On) cho nhiều ứng dụng nội bộ.')
doc.add_paragraph('Thành phần chính:', style='List Bullet')
doc.add_paragraph('AuthService: ASP.NET Core 10 Web API, quản lý User/Role/Auth, phát hành JWT.', style='List Bullet 2')
doc.add_paragraph('Application1: ASP.NET Core 10 MVC, dùng đăng nhập từ AuthService.', style='List Bullet 2')
doc.add_paragraph('Application2: ASP.NET Core 10 MVC, dùng đăng nhập từ AuthService.', style='List Bullet 2')
doc.add_paragraph('Cơ sở dữ liệu: SQL Server (Entity Framework Core).', style='List Bullet')

doc.add_heading('2. Kiến trúc triển khai', level=1)
doc.add_paragraph('Mô hình triển khai:', style='List Bullet')
doc.add_paragraph('AuthService chạy độc lập, cung cấp API xác thực + màn hình SSO login.', style='List Bullet 2')
doc.add_paragraph('App1/App2 xác thực bằng JWT middleware đọc access_token từ cookie.', style='List Bullet 2')
doc.add_paragraph('Cookie domain chung cho phép SSO giữa nhiều subdomain (môi trường production).', style='List Bullet 2')
doc.add_paragraph('Luồng chuẩn:', style='List Bullet')
doc.add_paragraph('User truy cập App1/App2 -> chưa có token -> redirect sang AuthService /sso/login.', style='List Bullet 2')
doc.add_paragraph('User đăng nhập thành công -> AuthService set cookie access_token + refresh_token.', style='List Bullet 2')
doc.add_paragraph('Redirect về app ban đầu -> app xác thực token hợp lệ và cho truy cập.', style='List Bullet 2')
doc.add_paragraph('User sang app còn lại -> cookie dùng lại -> không cần đăng nhập lại.', style='List Bullet 2')

doc.add_heading('3. Công nghệ sử dụng', level=1)
tech = doc.add_table(rows=1, cols=2)
tech.style = 'Table Grid'
tech.rows[0].cells[0].text = 'Thành phần'
tech.rows[0].cells[1].text = 'Công nghệ'
for k, v in [
    ('Backend', 'C# .NET 10'),
    ('API', 'ASP.NET Core Web API'),
    ('Web App', 'ASP.NET Core MVC + Razor'),
    ('ORM', 'Entity Framework Core'),
    ('Auth', 'JWT Bearer + Cookie HttpOnly Secure'),
    ('Password Hash', 'BCrypt'),
    ('Database', 'SQL Server'),
    ('Test API', 'Postman collection (đã xuất JSON)')
]:
    r = tech.add_row().cells
    r[0].text = k
    r[1].text = v

doc.add_heading('4. Thiết kế cơ sở dữ liệu (AuthService)', level=1)
for line in [
    'Users: Id, Username, PasswordHash, Email, FullName, IsActive, CreatedAt, RefreshToken, RefreshTokenExpiresAt',
    'Roles: Id, RoleName, Description',
    'UserRoles: Id, UserId, RoleId',
    'Permissions: Id, PermissionName, Description',
    'RolePermissions: Id, RoleId, PermissionId'
]:
    doc.add_paragraph(line, style='List Bullet')

doc.add_heading('5. Danh sách API chính (AuthService)', level=1)
for ep in [
    'POST /api/auth/login',
    'POST /api/auth/register',
    'POST /api/auth/refresh-token',
    'POST /api/auth/refresh-token-cookie',
    'POST /api/auth/logout',
    'GET /api/users',
    'POST /api/users',
    'PUT /api/users/{id}',
    'DELETE /api/users/{id}',
    'GET /api/roles',
    'POST /api/roles',
    'POST /api/roles/assign',
    'GET /sso/login',
    'POST /sso/login',
    'GET /sso/logout'
]:
    doc.add_paragraph(ep, style='List Bullet')

doc.add_paragraph('JWT claims chuẩn:', style='List Bullet')
for c in ['UserId', 'Username', 'Role']:
    doc.add_paragraph(c, style='List Bullet 2')

doc.add_heading('6. Chính sách bảo mật', level=1)
for item in [
    'Password hash bằng BCrypt.',
    'Access token hết hạn 60 phút.',
    'Refresh token hết hạn 7 ngày.',
    'Cookie access_token/refresh_token cấu hình HttpOnly + Secure + Path=/.',
    'Role-based authorization cho API quản trị người dùng/role.',
    'CORS dạng whitelist origin (không dùng *) cho App1/App2.',
    'Bắt buộc HTTPS trong luồng SSO.'
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('7. Cấu hình CORS', level=1)
doc.add_paragraph('AuthService áp policy StrictCors, chỉ cho phép các origin đã khai báo trong appsettings:', style='List Bullet')
doc.add_paragraph('https://localhost:7002', style='List Bullet 2')
doc.add_paragraph('https://localhost:7281', style='List Bullet 2')
doc.add_paragraph('Cho phép credentials để cookie SSO hoạt động qua trình duyệt.', style='List Bullet')

doc.add_heading('8. Cấu trúc solution thực tế', level=1)
for p in ['AuthService.slnx', 'AuthService/', 'Application1/', 'Application2/', 'docs/', 'postman/']:
    doc.add_paragraph(p, style='List Bullet')

doc.add_heading('9. Hướng dẫn cài đặt và chạy', level=1)
for step in [
    'Yêu cầu: .NET SDK 10, SQL Server, chứng chỉ dev HTTPS.',
    'Build: dotnet build AuthService.slnx',
    'Chạy AuthService: cd AuthService && dotnet run',
    'Chạy App1: cd Application1 && dotnet run',
    'Chạy App2: cd Application2 && dotnet run'
]:
    doc.add_paragraph(step, style='List Bullet')

doc.add_heading('10. Kiểm thử chức năng', level=1)
for t in [
    'Mở App1 -> redirect AuthService login -> đăng nhập thành công -> vào dashboard.',
    'Mở App2 -> truy cập trực tiếp không cần login lại.',
    'Logout SSO -> quay lại App1/App2 đều yêu cầu login lại.',
    'Regression tự động: powershell -ExecutionPolicy Bypass -File .\\sso-regression.ps1'
]:
    doc.add_paragraph(t, style='List Bullet')

doc.add_heading('11. Postman và tài liệu API', level=1)
doc.add_paragraph('Collection: postman/AuthService-SSO.postman_collection.json', style='List Bullet')
doc.add_paragraph('Environment: postman/AuthService-SSO.local.postman_environment.json', style='List Bullet')

doc.add_heading('12. Khuyến nghị triển khai production', level=1)
for rec in [
    'Đặt CookieDomain thành .company.com để SSO đa subdomain.',
    'Cập nhật AllowedReturnHosts theo domain thật (auth/app1/app2).',
    'Dùng chứng chỉ TLS hợp lệ.',
    'Đưa JWT key vào Secret Manager/Key Vault.',
    'Bật monitoring + auditing cho bảo mật.'
]:
    doc.add_paragraph(rec, style='List Bullet')

doc.add_heading('13. Lộ trình nâng cao', level=1)
doc.add_paragraph('Nếu mở rộng doanh nghiệp, cân nhắc OpenID Connect với IdentityServer hoặc Keycloak để quản trị phiên, federation và policy nâng cao.', style='List Bullet')

doc.add_paragraph('---')
doc.add_paragraph('Tài liệu được tạo tự động từ trạng thái workspace hiện tại.')

doc.save(file_path)
print(file_path)
